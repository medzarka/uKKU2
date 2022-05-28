import os
import shutil
import statistics
import random
import string
import zipfile
from io import StringIO  ## for Python 3
from shutil import copyfile
from shutil import make_archive
from wsgiref.util import FileWrapper
from pathlib import Path

from django.contrib.auth.models import User
from django.contrib import messages

from _control._Measurement.GradesFileTools import GradesFileReader

from django.core.files.storage import FileSystemStorage
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.urls import reverse

from _control._Measurement.Measurement_FS import Measurement_FS
from _control._Measurement.Measurement_tools import Course_Measurment, Department_Measurment, MEASUREMENT_Common
from _control._Measurement._masurement_Archieve_Maker import MeasurementArchiveMakerThread
from _control._Measurement._measurement_course_reports_admin import _measurement_course_reports_admin
from _control._Measurement._measurement_section_reports import _page_generate_section_reports
from _control._Measurement._measurement_course_reports import _page_generate_course_reports
from _control._Measurement._measurement_section_reports_admin import _measurement_section_reports_admin
from _control._Measurement._measurement_department_reports import _page_generate_department_reports
from _control._Measurement._measurement_reviewers import _page_measuerement_reviewers
from _control._quality._quality_Archieve_Maker import QualityArchiveMakerThread
from _control._quality._quality_mycfis import _page_mycfis
from _control._quality._quality_mycfis_admin import _page_mycfis_admin
from _control._quality._quality_mycfis_reviewers import _page_mycfis_reviewers
from _control._quality._quality_export import _page_quality_export
from _control._Measurement._measurement_export import _page_measuerement_export

from _control._dahsboard._dashboard_page import _dashboard
from _data._data_measurement import GradesFile, CourseFile, DepartmentFile, Department
from _data._data_periods import Semester
from _data._data_quality import Course_CFI

from django.http import HttpResponse
from docx import Document
from docx.shared import Inches

from django.http import HttpResponse
import xlsxwriter

from io import BytesIO


def change_semester(request):
    semester_id = request.POST['semester']
    request.session['selected_semester'] = int(semester_id)
    return redirect(reverse('dashboard'))


characters = list(string.ascii_letters + string.digits + "!@#$%&*")


def generate_random_password():
    length = 6
    random.shuffle(characters)
    password = []
    for i in range(length):
        password.append(random.choice(characters))
    random.shuffle(password)
    return "".join(password)


def reset_password(request):
    if request.method == 'POST':
        _email = request.POST.get('email', None)
        if _email is not None and _email != '':
            try:
                _user = User.objects.get(email=_email)
                _password = generate_random_password()
                _user.set_password(_password)
                _user.save()
                messages.success(request, f'The password is updated.')
                from _data._data_emails import email
                __email = email(email_receiver=_email
                                , email_title='Quality uKKU2 Password Update.'
                                , email_message=f'Your new password is {_password}. '
                                                f'You can login with the username {_user.username} .')
                __email.save()
                __email.send()
                messages.success(request,
                                 f'An email was sent containing the new password to {_email}. '
                                 f'Do not forget to check the SPAM folder.')
            except User.DoesNotExist:
                messages.error(request, 'User not find with provided Email ! ')
        else:
            messages.error(request, 'Invalid Email ! ')

    return render(request, "registration/reset_password.html", {})


@login_required
def dashboard(request):
    __page = _dashboard(request, 'dashboard')
    return render(request, "base.html", __page.getContext())


@login_required
def measurement_section_reports(request):
    __page = _page_generate_section_reports(request, 'measurement_section_reports')
    return render(request, "base.html", __page.getContext())


@login_required
def measurement_course_reports(request):
    __page = _page_generate_course_reports(request, 'measurement_course_reports')
    return render(request, "base.html", __page.getContext())


@login_required
def measurement_section_reports_admin(request):
    __page = _measurement_section_reports_admin(request, 'measurement_section_reports_admin')
    return render(request, "base.html", __page.getContext())


@login_required
def measurement_course_reports_admin(request):
    __page = _measurement_course_reports_admin(request, 'measurement_course_reports_admin')
    return render(request, "base.html", __page.getContext())


@login_required
def measurement_department_reports(request):
    __page = _page_generate_department_reports(request, 'measurement_department_reports')
    return render(request, "base.html", __page.getContext())


@login_required
def measurement__reviewers(request):
    __page = _page_measuerement_reviewers(request, 'measurement__reviewers')
    return render(request, "base.html", __page.getContext())


@login_required
def measurement_export(request):
    __page = _page_measuerement_export(request, 'measurement_export')
    return render(request, "base.html", __page.getContext())


@login_required
def quality_cfi_reports(request):
    __page = _page_mycfis(request, 'quality_mycfis')
    return render(request, "base.html", __page.getContext())


@login_required
def quality_mycfis_admin(request):
    __page = _page_mycfis_admin(request=request, link='quality_mycfis_admin')
    return render(request, "base.html", __page.getContext())


@login_required
def quality_mycfis_reviewers(request):
    __page = _page_mycfis_reviewers(request=request, link='quality_mycfis_reviewers')
    return render(request, "base.html", __page.getContext())


@login_required
def quality_export(request):
    __page = _page_quality_export(request=request, link='quality_export')
    return render(request, "base.html", __page.getContext())


def video(request):
    return render(request, "video.html", {})


def generate_section_report_list(request):
    print(request.POST.keys())
    if 'selected_semester' in request.POST.keys() and 'selected_action' in request.POST.keys():
        _selected_semester = Semester.objects.get(
            semester_id=int(request.POST['selected_semester']))

        document = Document()

        document.add_heading(
            'Section Report List for the semester ' + _selected_semester.semester_name + ' of the academic year ' + _selected_semester.semester_academic_year.academic_year_name,
            0)

        __departments = {}
        for _report in GradesFile.objects.filter(semester=_selected_semester):
            try:
                __departments[_report.section_department].append(_report)
            except KeyError:
                __departments[_report.section_department] = []
                __departments[_report.section_department].append(_report)

        for _department in __departments.keys():
            document.add_heading(_department, level=1)
            _nbr_rows = len(__departments[_department])
            table = document.add_table(rows=_nbr_rows + 1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'ID'
            hdr_cells[2].text = 'Campus'
            hdr_cells[3].text = 'Section'
            hdr_cells[4].text = 'Course'
            hdr_cells[5].text = 'Teacher'

            _i = 1
            for _report in __departments[_department]:
                hdr_cells = table.rows[_i].cells
                hdr_cells[0].text = str(_i)
                hdr_cells[1].text = str(_report.grades_file_id)
                hdr_cells[2].text = _report.campus_name
                hdr_cells[3].text = str(_report.section_code)
                hdr_cells[4].text = _report.course_name
                hdr_cells[5].text = _report.teacher.first_name
                _i += 1

        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=SectionList.docx'
        document.save(response)
        return response
    else:
        __page = _dashboard(request, 'dashboard')
        return render(request, "base.html", __page.getContext())


def generate_section_excel_list(request):
    if 'selected_semester' in request.POST.keys() and 'selected_action' in request.POST.keys():
        _selected_semester = Semester.objects.get(
            semester_id=int(request.POST['selected_semester']))
        print(1)

        response = HttpResponse(content_type='application/vnd.ms-excel')
        response['Content-Disposition'] = 'attachment; filename=Report.xlsx'
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output)
        worksheet = workbook.add_worksheet()
        print(2)

        worksheet.write(0, 0, 'Location')
        worksheet.write(0, 1, 'Department')
        worksheet.write(0, 2, 'Course Name')
        worksheet.write(0, 3, 'Course Code')
        worksheet.write(0, 4, 'Section Code')
        worksheet.write(0, 5, 'Mean')
        worksheet.write(0, 6, 'Standard Deviation')
        print(3)

        row = 1
        col = 0
        for _report in GradesFile.objects.filter(semester=_selected_semester):
            print(4)
            _campus = _report.campus_name
            _department = _report.section_department
            _course_name = _report.course_name
            _course_code = ''
            _section_code = str(_report.section_code)
            _mean = str(_report.stat_mean)
            _std = str(_report.stat_std)
            print(5)

            worksheet.write(row, col, _campus)
            worksheet.write(row, col + 1, _department)
            worksheet.write(row, col + 2, _course_name)
            worksheet.write(row, col + 3, _course_code)
            worksheet.write(row, col + 4, _section_code)
            worksheet.write(row, col + 5, _mean)
            worksheet.write(row, col + 6, _std)
            print(6)

            row += 1
            print(7)

        workbook.close()
        print(8)
        xlsx_data = output.getvalue()
        print(9)
        response.write(xlsx_data)
        print(10)
        return response
    else:
        __page = _dashboard(request, 'dashboard')
        return render(request, "base.html", __page.getContext())


def generate_course_report_list(request):
    if 'selected_semester' in request.POST.keys() and 'selected_action' in request.POST.keys():
        _selected_semester = Semester.objects.get(
            semester_id=int(request.POST['selected_semester']))

        document = Document()

        document.add_heading(
            'Course Report List for the semester ' + _selected_semester.semester_name + ' of the academic year ' + _selected_semester.semester_academic_year.academic_year_name,
            0)

        __departments = {}
        for _report in CourseFile.objects.filter(semester=_selected_semester):
            try:
                __departments[_report.course_department].append(_report)
            except KeyError:
                __departments[_report.course_department] = []
                __departments[_report.course_department].append(_report)

        for _department in __departments.keys():
            document.add_heading(_department, level=1)
            _nbr_rows = len(__departments[_department])
            table = document.add_table(rows=_nbr_rows + 1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'ID'
            hdr_cells[2].text = 'Campus'
            hdr_cells[3].text = 'Sections'
            hdr_cells[4].text = 'Course'
            hdr_cells[5].text = 'Teacher'

            _i = 1
            for _report in __departments[_department]:
                hdr_cells = table.rows[_i].cells
                hdr_cells[0].text = str(_i)
                hdr_cells[1].text = str(_report.course_file_id)
                hdr_cells[2].text = _report.campus_name
                hdr_cells[3].text = str(_report.section_codes)
                hdr_cells[4].text = _report.course_name
                hdr_cells[5].text = _report.teacher.first_name
                _i += 1

        document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=SectionList.docx'
        document.save(response)
        return response
    else:
        __page = _dashboard(request, 'dashboard')
        return render(request, "base.html", __page.getContext())


def generate_department_excel_list(request):
    if 'selected_semester' in request.POST.keys() and 'selected_action' in request.POST.keys():
        _selected_semester = Semester.objects.get(
            semester_id=int(request.POST['selected_semester']))

        ___departments = []
        ___courses = []
        ___files = []
        for _report in GradesFile.objects.filter(semester=_selected_semester):
            if _report.section_department not in ___departments:
                ___departments.append(_report.section_department)
            if _report.course_name not in ___courses:
                ___courses.append(_report.course_name)

        print(___departments)
        print(___courses)

        for __department in ___departments:
            from _control._Measurement.Measurement_FS import Measurement_FS
            _doc_filename = os.path.join('media/' + Measurement_FS.TMP.value,
                                         'Department_report_' + str(__department) + '.docx')
            print('Working with department ' + __department)
            _means = {}
            _grades = {}
            ___compuses = []
            _nbr_courses = 0
            for _course in ___courses:

                # print('\tWorking with the course ' + _course)
                __reports_for_course = GradesFile.objects.filter(semester=_selected_semester,
                                                                 section_department=__department, course_name=_course)
                for _report in __reports_for_course:
                    if _report.campus_name not in ___compuses:
                        ___compuses.append(_report.campus_name)

                _tool = Course_Measurment(_grades_files_objs=__reports_for_course, course_name=_course,
                                          course_file_obj=_course,
                                          department=__department)
                fused_grades, _ = _tool.extractGrades()
                if len(fused_grades['totals']) > 0:
                    _means[_course] = float("{0:.4f}".format(statistics.mean(fused_grades['totals'])))
                    _grades[_course] = fused_grades['totals']
                    _nbr_courses += 1

            _campuses_str = ''
            for _campus in ___compuses:
                if _campuses_str == '':
                    _campuses_str += _campus
                else:
                    _campuses_str += ', ' + _campus

            __tool2 = Department_Measurment(_grades, _means)
            _stat = __tool2.compute_statistics()
            _stat['department'] = __department
            _stat['campus'] = _campuses_str
            _stat['nbr_courses'] = str(_nbr_courses)
            print(_stat)
            __tool3 = MEASUREMENT_Common()
            __tool3.generate_department_docx_report(_stat, '', _doc_filename)
            print('----------')
            print('----------')
            print('----------')
            ___files.append(_doc_filename)

        from shutil import make_archive
        from wsgiref.util import FileWrapper
        files_path = os.path.join('media/' + Measurement_FS.TMP.value)
        path_to_zip = make_archive(files_path, "zip", files_path)
        response = HttpResponse(FileWrapper(open(path_to_zip, 'rb')), content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename="Docs.zip"'
        return response

    else:
        __page = _dashboard(request, 'dashboard')
        return render(request, "base.html", __page.getContext())


def generate_grades_excel_list(request):
    if 'selected_semester' in request.POST.keys() and 'selected_action' in request.POST.keys():
        _selected_semester = Semester.objects.get(
            semester_id=int(request.POST['selected_semester']))

        year_txt = _selected_semester.semester_academic_year.academic_year_name
        term_txt = _selected_semester.semester_name
        excel_filename = f'grades_list___{year_txt}__{term_txt}'

        response = HttpResponse(content_type='application/vnd.ms-excel')
        response['Content-Disposition'] = f'attachment; filename={excel_filename}.xlsx'
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output, {'nan_inf_to_errors': True})
        worksheet = workbook.add_worksheet()

        worksheet.write(0, 0, 'Campus')
        worksheet.write(0, 1, 'Department')
        worksheet.write(0, 2, 'Section')
        worksheet.write(0, 3, 'Course code')
        worksheet.write(0, 4, 'Course name')
        worksheet.write(0, 5, 'Student ID')
        worksheet.write(0, 6, 'Student Name')
        worksheet.write(0, 7, 'Mids')
        worksheet.write(0, 8, 'Final')
        worksheet.write(0, 9, 'Total')
        worksheet.write(0, 10, 'Grade')

        row_idx = 1
        col = 0
        for _report in GradesFile.objects.filter(semester=_selected_semester):
            _campus = _report.campus_name
            try:
                _department = _report.section_department.department_name
            except AttributeError:
                _department = ''
            _course_name = _report.course_name
            _course_code = _report.course_code
            _section_code = str(_report.section_code)

            _grades_filename = _report.grades_file.path
            tools = GradesFileReader(_grades_filename)
            tools.read()
            if tools.done == 1:
                grades = tools.data['grades']
                for row in grades:
                    worksheet.write(row_idx, col, _campus)
                    worksheet.write(row_idx, col + 1, _department)
                    worksheet.write(row_idx, col + 2, _section_code)
                    worksheet.write(row_idx, col + 3, _course_code)
                    worksheet.write(row_idx, col + 4, _course_name)
                    worksheet.write(row_idx, col + 5, row['student_id'])
                    worksheet.write(row_idx, col + 6, row['student_name'])
                    worksheet.write(row_idx, col + 7, row['mids'])
                    worksheet.write(row_idx, col + 8, row['finals'])
                    worksheet.write(row_idx, col + 9, row['totals'])
                    worksheet.write(row_idx, col + 10, row['grade'])
                    row_idx += 1

        workbook.close()
        xlsx_data = output.getvalue()
        response.write(xlsx_data)
        return response
    else:
        __page = _dashboard(request, 'dashboard')
        return render(request, "base.html", __page.getContext())


def make_archive(source, destination):
    base_name = '.'.join(destination.split('.')[:-1])
    format = destination.split('.')[-1]
    root_dir = os.path.dirname(source)
    base_dir = os.path.basename(source.strip(os.sep))
    shutil.make_archive(base_name, format, root_dir, base_dir)


def createDir(path):
    try:
        os.makedirs(path)
    except FileExistsError:
        pass
    except FileNotFoundError:
        pass


def generate_zipped_measurement_reports_list(request):
    if 'selected_semester' in request.POST.keys() and 'selected_action' in request.POST.keys():
        _selected_semester = Semester.objects.get(semester_id=int(request.POST['selected_semester']))
        _actual_user = User.objects.get(id=request.user.id)

        _thread = MeasurementArchiveMakerThread(_selected_semester, _actual_user)
        _thread.start()

    __page = _page_measuerement_export(request, 'measurement_export')
    return render(request, "base.html", __page.getContext())


def generate_zipped_quality_reports_list(request):
    if 'selected_semester' in request.POST.keys() and 'selected_action' in request.POST.keys():
        _selected_semester = Semester.objects.get(semester_id=int(request.POST['selected_semester']))
        _actual_user = User.objects.get(id=request.user.id)
        _selected_department = Department.objects.get(department_id=int(request.POST['department']))
        _selected_quality = request.POST['quality']

        _thread = QualityArchiveMakerThread(_selected_semester, _actual_user, _selected_department, _selected_quality,
                                            verbose=True)
        _thread.start()

    __page = _page_quality_export(request, 'quality_export')
    return render(request, "base.html", __page.getContext())
